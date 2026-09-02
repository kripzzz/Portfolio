import docx
import os

doc_path = os.path.join(os.path.dirname(__file__), 'KRUPA CRP (1).docx')
out_path = os.path.join(os.path.dirname(__file__), 'resume_extracted.txt')

doc = docx.Document(doc_path)
lines = []

lines.append("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        lines.append(f"{i}: {txt}")

lines.append("\n=== TABLES ===")
for t_idx, t in enumerate(doc.tables):
    lines.append(f"-- TABLE {t_idx+1} --")
    for r in t.rows:
        row_txt = " | ".join([c.text.strip().replace('\n', ' ') for c in r.cells])
        lines.append(row_txt)

with open(out_path, 'w', encoding='utf-8') as f:
    f.write("\n".join(lines))

print("Resume parsed successfully to resume_extracted.txt")
